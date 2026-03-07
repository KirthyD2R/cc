"""Convert PIPELINE_WORKFLOW.md and REQUIREMENTS_AND_FEATURES.md to Word documents."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elem)


def add_code_block(doc, text):
    """Add a code block with gray background."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles["No Spacing"]
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table_from_rows(doc, header_row, data_rows):
    """Add a formatted table."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = text.strip()
        set_cell_shading(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(data_rows):
        for c_idx, text in enumerate(row):
            if c_idx < cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = text.strip()
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def parse_table_block(lines):
    """Parse markdown table lines into header and data rows."""
    # Filter out separator lines (|---|---|)
    content_lines = [l for l in lines if not re.match(r"^\s*\|[\s\-:|]+\|\s*$", l)]
    rows = []
    for line in content_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 1:
        return rows[0], rows[1:]
    return [], []


def add_formatted_run(paragraph, text):
    """Add a run with inline formatting (bold, code, etc.)."""
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # Style headings
    for level in range(1, 5):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table block
        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, data = parse_table_block(table_lines)
            if header:
                add_table_from_rows(doc, header, data)
            continue

        # Bullet list
        bullet_match = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            level = min(indent // 2, 3)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            add_formatted_run(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if num_match:
            indent = len(num_match.group(1))
            text = num_match.group(2)
            level = min(indent // 3, 3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            add_formatted_run(p, text)
            i += 1
            continue

        # Sub-items with letters (a., b., etc.) or indented bullets
        sub_match = re.match(r"^\s+-\s+\*\*([a-h])\.\s+(.+)\*\*(.*)$", line)
        if sub_match:
            text = f"{sub_match.group(1)}. {sub_match.group(2)}{sub_match.group(3)}"
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_run(p, f"**{text}**")
            i += 1
            continue

        # Indented continuation (sub-bullets under numbered items)
        indent_bullet = re.match(r"^(\s{2,})[-*]\s+(.+)$", line)
        if indent_bullet:
            text = indent_bullet.group(2)
            indent_level = len(indent_bullet.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            add_formatted_run(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert_md_to_docx(
        os.path.join(DOCS_DIR, "PIPELINE_WORKFLOW.md"),
        os.path.join(DOCS_DIR, "PIPELINE_WORKFLOW.docx"),
    )
    convert_md_to_docx(
        os.path.join(DOCS_DIR, "REQUIREMENTS_AND_FEATURES.md"),
        os.path.join(DOCS_DIR, "REQUIREMENTS_AND_FEATURES.docx"),
    )
    print("Done. Both Word files generated.")
